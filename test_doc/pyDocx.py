print 'hello world'
print 'hello 2'

import docx

document = docx.Document()

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.save('demo.docx')